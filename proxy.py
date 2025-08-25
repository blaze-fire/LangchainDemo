# Advanced Tableau Test Case Generation System with Human Feedback Loop
# Using LangGraph, LangChain, and OpenAI with Functional Programming

import os
import json
import hashlib
import pickle
from typing import Dict, List, Any, TypedDict, Annotated, Sequence, Optional, Tuple
from enum import Enum
import operator
from datetime import datetime
import re
from pathlib import Path

# Core imports
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA

# LangGraph imports
from langgraph.graph import StateGraph, END
from langgraph.checkpoint import MemorySaver

# Document processing
import pandas as pd
import numpy as np
from docx import Document as DocxDocument
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment

# Async and caching
from functools import lru_cache, wraps
import asyncio
from concurrent.futures import ThreadPoolExecutor

# Logging
import logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ============================================
# Enhanced State Definition with Feedback Loop
# ============================================

class ConversationState(TypedDict):
    """Enhanced state with conversation history and feedback tracking"""
    messages: Annotated[Sequence[BaseMessage], operator.add]
    fsd_content: str
    fsd_chunks: List[Dict[str, Any]]
    t5_data: Dict[str, Any]
    disclosure_mapping: Dict[str, Any]
    regression_samples: List[str]
    extracted_requirements: Dict[str, Any]
    generated_test_cases: List[Dict[str, Any]]
    reviewed_test_cases: List[Dict[str, Any]]
    current_agent: str
    iteration: int
    user_feedback: List[Dict[str, Any]]
    feedback_history: List[Dict[str, Any]]
    modifications_made: List[str]
    satisfaction_score: float
    vectorstore: Optional[Any]
    test_case_versions: List[List[Dict[str, Any]]]
    metadata: Dict[str, Any]
    structured_data_insights: Dict[str, Any]
    quality_metrics: Dict[str, Any]

# ============================================
# Document Processing with Semantic Chunking
# ============================================

def create_semantic_chunks(content: str, chunk_size: int = 2000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
    """Create semantic chunks from large documents with metadata"""
    
    # Initialize text splitter with semantic awareness
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n## ", "\n\n### ", "\n\n", "\n", ". ", ", ", " "],
        length_function=len,
    )
    
    # Split text
    chunks = splitter.split_text(content)
    
    # Add metadata to each chunk
    chunk_dicts = []
    for i, chunk in enumerate(chunks):
        chunk_dict = {
            "content": chunk,
            "chunk_id": i,
            "chunk_size": len(chunk),
            "keywords": extract_keywords(chunk),
            "section_type": identify_section_type(chunk),
            "has_requirements": "must" in chunk.lower() or "shall" in chunk.lower() or "requirement" in chunk.lower(),
            "has_business_rules": "business rule" in chunk.lower() or "validation" in chunk.lower(),
            "has_calculations": "calculate" in chunk.lower() or "formula" in chunk.lower(),
        }
        chunk_dicts.append(chunk_dict)
    
    return chunk_dicts

def extract_keywords(text: str) -> List[str]:
    """Extract important keywords from text chunk"""
    # Common test-related keywords
    important_terms = [
        "validation", "calculation", "filter", "parameter", "dimension",
        "measure", "aggregate", "dashboard", "worksheet", "data source",
        "join", "blend", "extract", "refresh", "security", "permission",
        "performance", "user", "role", "access", "business rule"
    ]
    
    keywords = []
    text_lower = text.lower()
    for term in important_terms:
        if term in text_lower:
            keywords.append(term)
    
    return keywords

def identify_section_type(text: str) -> str:
    """Identify the type of section in the document"""
    text_lower = text.lower()
    
    if "requirement" in text_lower or "functional" in text_lower:
        return "requirements"
    elif "business rule" in text_lower or "validation" in text_lower:
        return "business_rules"
    elif "technical" in text_lower or "architecture" in text_lower:
        return "technical"
    elif "user interface" in text_lower or "ui" in text_lower:
        return "ui_ux"
    elif "data" in text_lower and "model" in text_lower:
        return "data_model"
    elif "test" in text_lower or "scenario" in text_lower:
        return "test_scenarios"
    else:
        return "general"

def process_fsd_document_advanced(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """Advanced FSD processing with semantic chunking and metadata extraction"""
    try:
        doc = DocxDocument(file_path)
        content = []
        document_structure = {"sections": [], "tables": [], "lists": []}
        
        # Extract all content with structure preservation
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check for headings
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name[-1] if paragraph.style.name[-1].isdigit() else '1'
                    content.append(f"{'#' * int(level)} {paragraph.text.strip()}")
                    document_structure["sections"].append({
                        "level": int(level),
                        "text": paragraph.text.strip()
                    })
                else:
                    content.append(paragraph.text.strip())
        
        # Extract and structure tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                
                if row_idx == 0:
                    headers = row_data
                    table_data.append("| " + " | ".join(row_data) + " |")
                    table_data.append("|" + "---|" * len(row_data))
                else:
                    table_data.append("| " + " | ".join(row_data) + " |")
            
            if table_data:
                table_content = "\n".join(table_data)
                content.append(f"\n### Table {table_idx + 1}\n{table_content}\n")
                document_structure["tables"].append({
                    "index": table_idx,
                    "headers": headers,
                    "row_count": len(table.rows)
                })
        
        full_content = "\n\n".join(content)
        
        # Create semantic chunks
        chunks = create_semantic_chunks(full_content)
        
        # Add document-level metadata
        for chunk in chunks:
            chunk["document_structure"] = document_structure
        
        return full_content, chunks
        
    except Exception as e:
        logger.error(f"Error processing FSD document: {str(e)}")
        return f"Error: {str(e)}", []

def process_excel_structured_data(file_path: str, file_type: str = "generic") -> Dict[str, Any]:
    """Process Excel files with careful structured data handling"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names
        
        structured_data = {
            "file_type": file_type,
            "sheets": {},
            "metadata": {
                "total_sheets": len(sheet_names),
                "processing_timestamp": datetime.now().isoformat()
            },
            "data_quality": {},
            "relationships": []
        }
        
        for sheet_name in sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Clean column names
            df.columns = df.columns.str.strip().str.replace(' ', '_')
            
            # Analyze data quality
            quality_metrics = analyze_data_quality(df)
            
            # Detect data types and patterns
            column_analysis = {}
            for col in df.columns:
                column_analysis[col] = {
                    "dtype": str(df[col].dtype),
                    "unique_values": int(df[col].nunique()),
                    "null_count": int(df[col].isnull().sum()),
                    "sample_values": df[col].dropna().head(3).tolist() if len(df[col].dropna()) > 0 else [],
                    "patterns": detect_patterns(df[col])
                }
            
            # Store structured information
            structured_data["sheets"][sheet_name] = {
                "shape": df.shape,
                "columns": df.columns.tolist(),
                "column_analysis": column_analysis,
                "data": df.to_dict('records'),
                "summary_stats": df.describe(include='all').to_dict() if len(df) > 0 else {},
                "quality_metrics": quality_metrics
            }
            
            structured_data["data_quality"][sheet_name] = quality_metrics
            
            # Detect potential relationships between sheets
            if len(structured_data["sheets"]) > 1:
                relationships = detect_sheet_relationships(structured_data["sheets"])
                structured_data["relationships"] = relationships
        
        return structured_data
        
    except Exception as e:
        logger.error(f"Error processing Excel file: {str(e)}")
        return {"error": str(e)}

def analyze_data_quality(df: pd.DataFrame) -> Dict[str, Any]:
    """Analyze data quality metrics for structured data"""
    metrics = {
        "completeness": 1 - (df.isnull().sum().sum() / (df.shape[0] * df.shape[1])) if df.shape[0] > 0 else 0,
        "duplicate_rows": df.duplicated().sum(),
        "duplicate_percentage": (df.duplicated().sum() / len(df) * 100) if len(df) > 0 else 0,
        "column_metrics": {}
    }
    
    for col in df.columns:
        metrics["column_metrics"][col] = {
            "completeness": 1 - (df[col].isnull().sum() / len(df)) if len(df) > 0 else 0,
            "unique_ratio": df[col].nunique() / len(df) if len(df) > 0 else 0,
            "most_frequent": df[col].mode().iloc[0] if len(df[col].mode()) > 0 else None
        }
    
    return metrics

# ============================================
# Human Feedback Loop Implementation
# ============================================

def create_feedback_analyzer(llm: ChatOpenAI):
    """Analyze user feedback and determine required modifications"""
    
    def analyze_feedback(state: ConversationState) -> Dict:
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert at understanding user feedback for test cases.
            
            Analyze the user's feedback and determine:
            1. INTENT: What changes are requested?
               - Add new test cases
               - Modify existing test cases
               - Remove test cases
               - Change priorities
               - Improve coverage
               - Fix issues
            
            2. SCOPE: Which test cases are affected?
               - Specific test IDs
               - Categories
               - All test cases
            
            3. DETAILS: What specific changes to make?
            
            4. SENTIMENT: User satisfaction level (1-10)
            
            Output as JSON:
            {
                "intent": "modify/add/remove/improve",
                "scope": {"test_ids": [...], "categories": [...]},
                "specific_changes": [...],
                "sentiment_score": 7,
                "clarification_needed": false,
                "clarification_questions": []
            }"""),
            ("human", """Current Test Cases Summary:
            {test_summary}
            
            User Feedback:
            {feedback}
            
            Analyze this feedback and determine required actions:""")
        ])
        
        # Prepare test summary
        test_summary = summarize_test_cases(state["generated_test_cases"])
        latest_feedback = state["user_feedback"][-1] if state["user_feedback"] else {"feedback": "No feedback"}
        
        chain = prompt | llm
        response = chain.invoke({
            "test_summary": test_summary,
            "feedback": latest_feedback.get("feedback", "")
        })
        
        try:
            feedback_analysis = json.loads(response.content if hasattr(response, 'content') else str(response))
        except:
            feedback_analysis = {
                "intent": "modify",
                "scope": {"test_ids": [], "categories": []},
                "specific_changes": [response.content if hasattr(response, 'content') else str(response)],
                "sentiment_score": 5,
                "clarification_needed": False
            }
        
        state["feedback_history"].append({
            "timestamp": datetime.now().isoformat(),
            "analysis": feedback_analysis,
            "original_feedback": latest_feedback
        })
        
        state["satisfaction_score"] = feedback_analysis.get("sentiment_score", 5)
        
        return state
    
    return analyze_feedback

def create_test_modifier(llm: ChatOpenAI):
    """Modify test cases based on analyzed feedback"""
    
    def modify_test_cases(state: ConversationState) -> Dict:
        latest_analysis = state["feedback_history"][-1]["analysis"] if state["feedback_history"] else {}
        intent = latest_analysis.get("intent", "modify")
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", f"""You are modifying test cases based on user feedback.
            
            MODIFICATION INTENT: {intent}
            
            Rules:
            1. Preserve test cases that weren't mentioned in feedback
            2. Apply changes precisely as requested
            3. Maintain test case structure and quality
            4. Update version tracking
            5. Document all changes made
            
            For ADDITIONS:
            - Create new test cases meeting the requirements
            - Ensure unique IDs
            - Follow established format
            
            For MODIFICATIONS:
            - Update only specified aspects
            - Preserve other fields
            - Maintain relationships
            
            For REMOVALS:
            - Remove specified test cases
            - Update related test dependencies
            
            Output the complete updated test case list."""),
            ("human", """Current Test Cases:
            {current_tests}
            
            Required Changes:
            {changes}
            
            Original Requirements (for reference):
            {requirements}
            
            Apply the modifications and return updated test cases:""")
        ])
        
        chain = prompt | llm
        response = chain.invoke({
            "current_tests": json.dumps(state["generated_test_cases"], indent=2)[:8000],
            "changes": json.dumps(latest_analysis.get("specific_changes", []), indent=2),
            "requirements": json.dumps(state["extracted_requirements"], indent=2)[:3000]
        })
        
        # Parse modified test cases
        modified_tests = parse_test_cases_advanced(response.content if hasattr(response, 'content') else str(response))
        
        # Track modifications
        modifications = []
        for new_tc in modified_tests:
            # Find if this is a modified version of existing test
            existing = next((tc for tc in state["generated_test_cases"] 
                           if tc.get("tc_id") == new_tc.get("tc_id")), None)
            if existing:
                changes = compare_test_cases(existing, new_tc)
                if changes:
                    modifications.append({
                        "tc_id": new_tc.get("tc_id"),
                        "changes": changes
                    })
            else:
                modifications.append({
                    "tc_id": new_tc.get("tc_id"),
                    "changes": ["New test case added"]
                })
        
        state["generated_test_cases"] = modified_tests
        state["test_case_versions"].append(modified_tests.copy())
        state["modifications_made"] = modifications
        state["iteration"] += 1
        
        # Recalculate quality metrics
        state["quality_metrics"] = calculate_quality_metrics(modified_tests, state["extracted_requirements"])
        
        state["messages"].append(AIMessage(content=f"Modified test cases based on feedback: {len(modifications)} changes made"))
        
        return state
    
    return modify_test_cases

def summarize_test_cases(test_cases: List[Dict]) -> str:
    """Create a summary of test cases for feedback analysis"""
    summary = f"Total Test Cases: {len(test_cases)}\n\n"
    
    # Group by category
    categories = {}
    for tc in test_cases:
        cat = tc.get("category", "Uncategorized")
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(tc)
    
    for cat, tcs in categories.items():
        summary += f"{cat}: {len(tcs)} test cases\n"
        for tc in tcs[:3]:  # Show first 3 as examples
            summary += f"  - {tc.get('tc_id', 'No ID')}: {tc.get('name', 'No name')[:50]}\n"
        if len(tcs) > 3:
            summary += f"  ... and {len(tcs) - 3} more\n"
    
    return summary

def compare_test_cases(old_tc: Dict, new_tc: Dict) -> List[str]:
    """Compare two test cases and return list of changes"""
    changes = []
    
    for field in ["name", "priority", "category", "steps", "expected_results", "test_data"]:
        if old_tc.get(field) != new_tc.get(field):
            changes.append(f"Updated {field}")
    
    return changes

# ============================================
# Interactive Session Manager
# ============================================

def create_interactive_session(llm: ChatOpenAI):
    """Manage interactive session with user"""
    
    def interactive_handler(state: ConversationState) -> Dict:
        """Handle interactive session with user feedback"""
        
        # Check if this is first iteration or subsequent
        if state["iteration"] == 0:
            message = format_initial_test_cases(state)
        else:
            message = format_modification_results(state)
        
        state["messages"].append(AIMessage(content=message))
        
        # Determine if we should continue or end
        if state["satisfaction_score"] >= 8 or state["iteration"] >= 5:
            state["current_agent"] = "finalize"
        else:
            state["current_agent"] = "await_feedback"
        
        return state
    
    return interactive_handler

def format_initial_test_cases(state: ConversationState) -> str:
    """Format initial test cases for user review"""
    metrics = state.get("quality_metrics", {})
    test_cases = state.get("generated_test_cases", [])
    
    message = f"""
## Test Case Generation Complete! 📋

### Summary Statistics:
- **Total Test Cases Generated**: {metrics.get('total_test_cases', 0)}
- **Requirement Coverage**: {metrics.get('completeness_score', 0):.1f}%
- **Quality Score**: {metrics.get('quality_score', 0):.1f}%

### Priority Distribution:
"""
    
    for priority, count in metrics.get('priority_distribution', {}).items():
        message += f"- {priority}: {count} test cases\n"
    
    message += "\n### Category Distribution:\n"
    for category, count in metrics.get('category_distribution', {}).items():
        message += f"- {category}: {count} test cases\n"
    
    message += "\n### Sample Test Cases:\n"
    for tc in test_cases[:3]:
        message += f"""
**{tc.get('tc_id', 'No ID')}**: {tc.get('name', 'Unnamed')}
- Priority: {tc.get('priority', 'Not set')}
- Category: {tc.get('category', 'Not set')}
"""
    
    message += """

### 📝 Your Feedback Options:
1. **Approve**: "Looks good" / "Approve" / "Generate Excel"
2. **Modify Specific Tests**: "Change TC_001 priority to Critical"
3. **Add Tests**: "Add more security test cases"
4. **Remove Tests**: "Remove all low priority tests"
5. **Improve Coverage**: "Need more data validation tests"

**Please provide your feedback:**"""
    
    return message

def format_modification_results(state: ConversationState) -> str:
    """Format modification results for user review"""
    modifications = state.get("modifications_made", [])
    metrics = state.get("quality_metrics", {})
    
    message = f"""
## ✅ Test Cases Updated Based on Your Feedback

### Changes Made:
"""
    
    for mod in modifications[:5]:
        message += f"- {mod['tc_id']}: {', '.join(mod['changes'])}\n"
    
    if len(modifications) > 5:
        message += f"... and {len(modifications) - 5} more modifications\n"
    
    message += f"""

### Updated Metrics:
- **Total Test Cases**: {metrics.get('total_test_cases', 0)}
- **Requirement Coverage**: {metrics.get('completeness_score', 0):.1f}%
- **Quality Score**: {metrics.get('quality_score', 0):.1f}%

### What would you like to do next?
1. **Approve and Generate Excel**: "Approved" / "Generate Excel"
2. **Further Modifications**: Describe additional changes needed
3. **View Specific Tests**: "Show me security test cases"

**Your feedback:**"""
    
    return message

# ============================================
# Export Functions with Formatting
# ============================================

def export_test_cases_to_excel_formatted(test_cases: List[Dict], 
                                        output_path: str,
                                        metadata: Dict = None,
                                        quality_metrics: Dict = None):
    """Export test cases to formatted Excel with multiple sheets"""
    
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Main test cases sheet
        df_tests = pd.DataFrame(test_cases)
        
        # Reorder columns for better readability
        column_order = ["tc_id", "name", "priority", "category", "objective", 
                       "prerequisites", "steps", "test_data", "expected_results", "notes"]
        existing_cols = [col for col in column_order if col in df_tests.columns]
        other_cols = [col for col in df_tests.columns if col not in column_order]
        df_tests = df_tests[existing_cols + other_cols]
        
        df_tests.to_excel(writer, sheet_name='Test Cases', index=False)
        
        # Format the test cases sheet
        worksheet = writer.sheets['Test Cases']
        
        # Add filters
        worksheet.auto_filter.ref = worksheet.dimensions
        
        # Adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Add summary sheet
        if quality_metrics:
            summary_data = {
                "Metric": ["Total Test Cases", "Requirement Coverage", "Quality Score", 
                          "Critical Tests", "High Priority Tests", "Medium Priority Tests", "Low Priority Tests"],
                "Value": [
                    quality_metrics.get("total_test_cases", 0),
                    f"{quality_metrics.get('completeness_score', 0):.1f}%",
                    f"{quality_metrics.get('quality_score', 0):.1f}%",
                    quality_metrics.get("priority_distribution", {}).get("Critical", 0),
                    quality_metrics.get("priority_distribution", {}).get("High", 0),
                    quality_metrics.get("priority_distribution", {}).get("Medium", 0),
                    quality_metrics.get("priority_distribution", {}).get("Low", 0)
                ]
            }
            df_summary = pd.DataFrame(summary_data)
            df_summary.to_excel(writer, sheet_name='Summary', index=False)
        
        # Add requirement coverage sheet
        if quality_metrics and "requirement_coverage" in quality_metrics:
            coverage_data = []
            for req_type, coverage in quality_metrics["requirement_coverage"].items():
                coverage_data.append({
                    "Requirement Type": req_type,
                    "Total Requirements": coverage.get("total", 0),
                    "Covered": coverage.get("covered", 0),
                    "Coverage %": f"{coverage.get('percentage', 0):.1f}%"
                })
            
            if coverage_data:
                df_coverage = pd.DataFrame(coverage_data)
                df_coverage.to_excel(writer, sheet_name='Coverage', index=False)
    
    logger.info(f"Test cases exported to {output_path}")
    return output_path

# ============================================
# Graph Construction with Feedback Loop
# ============================================

def build_interactive_test_generation_graph(llm: ChatOpenAI, embeddings: OpenAIEmbeddings):
    """Build LangGraph workflow with human feedback loop"""
    
    # Create vector store (will be populated during execution)
    vectorstore = None
    
    # Create agents
    def create_agents():
        return {
            "fsd_analyzer": create_advanced_fsd_analyzer(llm, vectorstore) if vectorstore else None,
            "structured_analyzer": create_structured_data_analyzer(llm),
            "test_generator": create_intelligent_test_generator(llm),
            "feedback_analyzer": create_feedback_analyzer(llm),
            "test_modifier": create_test_modifier(llm),
            "interactive_handler": create_interactive_session(llm)
        }
    
    # Define conditional edge logic
    def should_continue(state: ConversationState) -> str:
        """Determine next step based on state"""
        
        if state.get("current_agent") == "await_feedback":
            # Wait for user feedback
            return "feedback_wait"
        elif state.get("current_agent") == "finalize":
            return "export"
        elif state["iteration"] == 0:
            # First iteration - generate initial tests
            return "generate"
        else:
            # Subsequent iterations - modify based on feedback
            return "modify"
    
    # Create workflow
    workflow = StateGraph(ConversationState)
    
    # Add nodes
    workflow.add_node("analyze_fsd", lambda state: state)  # Placeholder, will be replaced
    workflow.add_node("analyze_structured", lambda state: state)  # Placeholder
    workflow.add_node("generate_tests", lambda state: state)  # Placeholder
    workflow.add_node("analyze_feedback", lambda state: state)  # Placeholder
    workflow.add_node("modify_tests", lambda state: state)  # Placeholder
    workflow.add_node("interactive", lambda state: state)  # Placeholder
    workflow.add_node("export", lambda state: state)  # Final node
    
    # Define edges
    workflow.set_entry_point("analyze_fsd")
    workflow.add_edge("analyze_fsd", "analyze_structured")
    workflow.add_edge("analyze_structured", "generate_tests")
    workflow.add_edge("generate_tests", "interactive")
    
    # Conditional edges for feedback loop
    workflow.add_conditional_edges(
        "interactive",
        should_continue,
        {
            "feedback_wait": "analyze_feedback",
            "export": "export",
            "generate": "generate_tests",
            "modify": "modify_tests"
        }
    )
    
    workflow.add_edge("analyze_feedback", "modify_tests")
    workflow.add_edge("modify_tests", "interactive")
    workflow.add_edge("export", END)
    
    # Compile
    memory = MemorySaver()
    app = workflow.compile(checkpointer=memory)
    
    return app, create_agents

# ============================================
# Main Execution Function with Feedback Loop
# ============================================

def generate_tableau_test_cases_interactive(
    fsd_path: str,
    t5_path: str,
    mapping_path: str,
    regression_path: str,
    openai_api_key: str,
    model: str = "gpt-4-turbo-preview",
    max_iterations: int = 5
) -> Dict[str, Any]:
    """Main function with interactive feedback loop"""
    
    # Initialize LLM and embeddings
    llm = ChatOpenAI(
        api_key=openai_api_key,
        model=model,
        temperature=0.3
    )
    
    embeddings = OpenAIEmbeddings(api_key=openai_api_key)
    
    # Process input documents
    print("\n📁 Processing input documents...")
    fsd_content, fsd_chunks = process_fsd_document_advanced(fsd_path)
    t5_data = process_excel_structured_data(t5_path, "t5_extract")
    disclosure_mapping = process_excel_structured_data(mapping_path, "disclosure_mapping")
    
    # Process regression minimally
    regression_samples = []
    try:
        df_regression = pd.read_excel(regression_path)
        if 'Test_Case' in df_regression.columns:
            regression_samples = ["Format: ID | Name | Steps | Expected"]
    except:
        pass
    
    print(f"✅ Processed: {len(fsd_chunks)} FSD chunks, {len(t5_data.get('sheets', {}))} T5 sheets")
    
    # Create vector store for RAG
    print("\n🔍 Creating knowledge base...")
    all_data = {
        "sheets": {**t5_data.get("sheets", {}), **disclosure_mapping.get("sheets", {})}
    }
    vectorstore = create_vector_store(fsd_chunks, all_data, embeddings)
    
    # Build graph
    print("\n🔧 Building intelligent workflow...")
    app, create_agents_func = build_interactive_test_generation_graph(llm, embeddings)
    
    # Create agents with vectorstore
    agents = create_agents_func()
    
    # Update nodes with actual agents
    for node_name, agent_func in agents.items():
        if agent_func:
            app.nodes[node_name.replace("_", "")] = agent_func
    
    # Prepare initial state
    initial_state = {
        "messages": [HumanMessage(content="Starting intelligent test case generation")],
        "fsd_content": fsd_content,
        "fsd_chunks": fsd_chunks,
        "t5_data": t5_data,
        "disclosure_mapping": disclosure_mapping,
        "regression_samples": regression_samples,
        "extracted_requirements": {},
        "generated_test_cases": [],
        "reviewed_test_cases": [],
        "current_agent": "analyze_fsd",
        "iteration": 0,
        "user_feedback": [],
        "feedback_history": [],
        "modifications_made": [],
        "satisfaction_score": 0,
        "vectorstore": vectorstore,
        "test_case_versions": [],
        "metadata": {
            "generation_date": datetime.now().isoformat(),
            "model": model,
            "fsd_path": fsd_path
        },
        "structured_data_insights": {},
        "quality_metrics": {}
    }
    
    # Run initial generation
    print("\n🚀 Generating initial test cases...")
    config = {"configurable": {"thread_id": f"test_gen_{datetime.now().strftime('%Y%m%d_%H%M%S')}"}}
    
    # Execute first pass
    state = initial_state
    state = agents["fsd_analyzer"](state) if agents["fsd_analyzer"] else state
    state = agents["structured_analyzer"](state)
    state = agents["test_generator"](state)
    state = agents["interactive_handler"](state)
    
    # Display initial results
    print("\n" + "="*60)
    print(state["messages"][-1].content)
    print("="*60)
    
    # Interactive feedback loop
    iteration = 0
    while iteration < max_iterations and state.get("current_agent") != "finalize":
        user_input = input("\n👤 Your feedback (or 'approve' to finish): ").strip()
        
        if user_input.lower() in ["approve", "approved", "done", "finish", "good", "looks good", "generate excel"]:
            state["satisfaction_score"] = 10
            state["current_agent"] = "finalize"
            break
        
        # Add user feedback
        state["user_feedback"].append({
            "feedback": user_input,
            "timestamp": datetime.now().isoformat()
        })
        
        # Analyze and apply feedback
        print("\n🔄 Processing your feedback...")
        state = agents["feedback_analyzer"](state)
        state = agents["test_modifier"](state)
        state = agents["interactive_handler"](state)
        
        # Display results
        print("\n" + "="*60)
        print(state["messages"][-1].content)
        print("="*60)
        
        iteration += 1
    
    # Export final results
    print("\n📊 Generating final Excel report...")
    output_path = f"test_cases_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    export_test_cases_to_excel_formatted(
        state["generated_test_cases"],
        output_path,
        state["metadata"],
        state["quality_metrics"]
    )
    
    print(f"\n✅ Success! Test cases exported to: {output_path}")
    print(f"\nFinal Statistics:")
    print(f"- Total Test Cases: {len(state['generated_test_cases'])}")
    print(f"- Iterations: {state['iteration']}")
    print(f"- Quality Score: {state['quality_metrics'].get('quality_score', 0):.1f}%")
    print(f"- Coverage: {state['quality_metrics'].get('completeness_score', 0):.1f}%")
    
    return {
        "test_cases": state["generated_test_cases"],
        "output_file": output_path,
        "metrics": state["quality_metrics"],
        "feedback_history": state["feedback_history"],
        "versions": state["test_case_versions"]
    }

# ============================================
# Main Entry Point
# ============================================

def main():
    """Main entry point with example usage"""
    
    # Configuration
    config = {
        "fsd_path": "path/to/fsd_document.docx",
        "t5_path": "path/to/t5_extract.xlsx",
        "mapping_path": "path/to/disclosure_mapping.xlsx",
        "regression_path": "path/to/regression_tests.xlsx",
        "openai_api_key": os.getenv("OPENAI_API_KEY", "your-api-key"),
        "model": "gpt-4-turbo-preview",
        "max_iterations": 5
    }
    
    try:
        # Run interactive generation
        results = generate_tableau_test_cases_interactive(**config)
        
        print("\n" + "="*60)
        print("🎉 Test Case Generation Complete!")
        print("="*60)
        print(f"\nOutput saved to: {results['output_file']}")
        print(f"Total test cases: {len(results['test_cases'])}")
        
        return results
        
    except Exception as e:
        logger.error(f"Error in test generation: {str(e)}")
        raise

if __name__ == "__main__":
    results = main()

def detect_patterns(series: pd.Series) -> Dict[str, Any]:
    """Detect patterns in data series"""
    patterns = {
        "is_numeric": pd.api.types.is_numeric_dtype(series),
        "is_datetime": pd.api.types.is_datetime64_any_dtype(series),
        "is_categorical": series.nunique() < len(series) * 0.5 if len(series) > 0 else False,
        "has_pattern": False,
        "pattern_type": None
    }
    
    # Check for common patterns
    if patterns["is_numeric"]:
        if series.min() >= 0 and series.max() <= 100:
            patterns["pattern_type"] = "percentage"
        elif all(series.dropna() == series.dropna().astype(int)):
            patterns["pattern_type"] = "integer"
    
    return patterns

def detect_sheet_relationships(sheets: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Detect potential relationships between sheets based on column names"""
    relationships = []
    sheet_names = list(sheets.keys())
    
    for i, sheet1 in enumerate(sheet_names):
        for sheet2 in sheet_names[i+1:]:
            cols1 = set(sheets[sheet1]["columns"])
            cols2 = set(sheets[sheet2]["columns"])
            
            common_columns = cols1.intersection(cols2)
            if common_columns:
                relationships.append({
                    "sheet1": sheet1,
                    "sheet2": sheet2,
                    "common_columns": list(common_columns),
                    "relationship_strength": len(common_columns) / min(len(cols1), len(cols2))
                })
    
    return relationships

# ============================================
# Vector Store Creation for RAG
# ============================================

def create_vector_store(fsd_chunks: List[Dict[str, Any]], 
                        structured_data: Dict[str, Any],
                        embeddings_model: OpenAIEmbeddings) -> Chroma:
    """Create vector store from document chunks for RAG"""
    
    texts = []
    metadatas = []
    
    # Add FSD chunks
    for chunk in fsd_chunks:
        texts.append(chunk["content"])
        metadatas.append({
            "source": "fsd",
            "chunk_id": chunk["chunk_id"],
            "section_type": chunk["section_type"],
            "has_requirements": chunk["has_requirements"]
        })
    
    # Add structured data as searchable text
    for sheet_name, sheet_data in structured_data.get("sheets", {}).items():
        # Create searchable representation of structured data
        sheet_text = f"Sheet: {sheet_name}\n"
        sheet_text += f"Columns: {', '.join(sheet_data['columns'])}\n"
        sheet_text += f"Data Quality: {json.dumps(sheet_data.get('quality_metrics', {}))}\n"
        
        texts.append(sheet_text)
        metadatas.append({
            "source": "structured_data",
            "sheet_name": sheet_name,
            "data_type": structured_data.get("file_type", "unknown")
        })
    
    # Create vector store
    vectorstore = Chroma.from_texts(
        texts=texts,
        metadatas=metadatas,
        embedding=embeddings_model,
        collection_name="test_generation_rag"
    )
    
    return vectorstore

# ============================================
# Enhanced Agent Functions with Quality Improvements
# ============================================

def create_advanced_fsd_analyzer(llm: ChatOpenAI, vectorstore: Chroma):
    """Advanced FSD analyzer with RAG and pattern recognition"""
    
    def analyze_fsd_advanced(state: ConversationState) -> Dict:
        # Use RAG to find relevant chunks
        retriever = vectorstore.as_retriever(
            search_kwargs={"k": 5, "filter": {"source": "fsd"}}
        )
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert FSD analyst specializing in Tableau test requirements.
            Analyze the FSD content with extreme attention to detail.
            
            EXTRACTION PRIORITIES:
            1. Mandatory Requirements (MUST test):
               - Regulatory compliance items
               - Critical business rules
               - Data integrity constraints
               - Security requirements
            
            2. Functional Requirements:
               - User workflows and scenarios
               - Data transformations
               - Calculations and formulas
               - Report layouts and formatting
            
            3. Non-Functional Requirements:
               - Performance benchmarks
               - Scalability requirements
               - Usability standards
               - Accessibility requirements
            
            4. Integration Points:
               - Data source connections
               - API interactions
               - Export/Import functionality
            
            5. Edge Cases and Boundaries:
               - Data limits
               - Null/empty handling
               - Error scenarios
               
            IMPORTANT: Extract EVERY testable item. Missing requirements means incomplete test coverage.
            
            Structure your output as:
            {
                "critical_requirements": [...],
                "functional_requirements": [...],
                "non_functional_requirements": [...],
                "integration_requirements": [...],
                "edge_cases": [...],
                "data_validations": [...],
                "user_scenarios": [...]
            }"""),
            ("human", """FSD Content Chunks:
            {fsd_chunks}
            
            Document Structure:
            {document_structure}
            
            Extract ALL testable requirements with their priorities and dependencies:""")
        ])
        
        # Get most relevant chunks
        relevant_chunks = retriever.get_relevant_documents("requirements validation calculations")
        chunk_texts = "\n---\n".join([doc.page_content for doc in relevant_chunks])
        
        # Extract document structure
        doc_structure = state["fsd_chunks"][0]["document_structure"] if state["fsd_chunks"] else {}
        
        chain = prompt | llm
        response = chain.invoke({
            "fsd_chunks": chunk_texts,
            "document_structure": json.dumps(doc_structure, indent=2)
        })
        
        # Parse and structure requirements
        try:
            requirements = json.loads(response.content if hasattr(response, 'content') else str(response))
        except:
            # Fallback parsing if JSON fails
            requirements = parse_requirements_from_text(response.content if hasattr(response, 'content') else str(response))
        
        state["extracted_requirements"] = requirements
        state["messages"].append(AIMessage(content=f"Extracted comprehensive requirements from FSD"))
        
        # Calculate requirement metrics
        total_reqs = sum(len(v) if isinstance(v, list) else 1 for v in requirements.values())
        state["metadata"]["total_requirements"] = total_reqs
        
        return state
    
    return analyze_fsd_advanced

def parse_requirements_from_text(text: str) -> Dict[str, List[str]]:
    """Fallback parser for requirements extraction"""
    requirements = {
        "critical_requirements": [],
        "functional_requirements": [],
        "non_functional_requirements": [],
        "integration_requirements": [],
        "edge_cases": [],
        "data_validations": [],
        "user_scenarios": []
    }
    
    current_category = None
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Check for category headers
        for category in requirements.keys():
            if category.replace('_', ' ').lower() in line.lower():
                current_category = category
                break
        
        # Add line to current category if it's a requirement
        if current_category and line and not any(cat in line.lower() for cat in requirements.keys()):
            if line.startswith(('-', '*', '•', '1', '2', '3', '4', '5', '6', '7', '8', '9')):
                requirements[current_category].append(line.lstrip('-*•0123456789. '))
    
    return requirements

def create_structured_data_analyzer(llm: ChatOpenAI):
    """Analyze structured Excel data for test scenarios"""
    
    def analyze_structured_data(state: ConversationState) -> Dict:
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are a data quality and testing expert analyzing structured data.
            
            Based on the structured data analysis, identify:
            
            1. DATA QUALITY TEST SCENARIOS:
               - Completeness checks for each critical field
               - Data type validations
               - Range and boundary validations
               - Referential integrity checks
               - Duplicate detection scenarios
            
            2. BUSINESS LOGIC VALIDATIONS:
               - Cross-field validations
               - Calculated field verifications
               - Aggregation accuracy tests
               - Data transformation validations
            
            3. PERFORMANCE TEST SCENARIOS:
               - Large dataset handling
               - Query optimization tests
               - Refresh rate validations
               - Concurrent user scenarios
            
            4. EDGE CASES FROM DATA PATTERNS:
               - Null/empty value handling
               - Special characters in text fields
               - Boundary values for numeric fields
               - Date range validations
            
            Provide specific, actionable test scenarios based on actual data characteristics."""),
            ("human", """Structured Data Analysis:
            
            T5 Extract:
            {t5_data}
            
            Disclosure Mapping:
            {disclosure_mapping}
            
            Data Quality Metrics:
            {quality_metrics}
            
            Generate comprehensive data-driven test scenarios:""")
        ])
        
        # Prepare structured data insights
        t5_summary = json.dumps(state["t5_data"], indent=2)[:4000]
        mapping_summary = json.dumps(state["disclosure_mapping"], indent=2)[:4000]
        
        quality_metrics = {
            "t5_quality": state["t5_data"].get("data_quality", {}),
            "mapping_quality": state["disclosure_mapping"].get("data_quality", {})
        }
        
        chain = prompt | llm
        response = chain.invoke({
            "t5_data": t5_summary,
            "disclosure_mapping": mapping_summary,
            "quality_metrics": json.dumps(quality_metrics, indent=2)
        })
        
        # Store structured data insights
        state["structured_data_insights"] = {
            "test_scenarios": response.content if hasattr(response, 'content') else str(response),
            "data_quality_issues": identify_quality_issues(state["t5_data"], state["disclosure_mapping"]),
            "recommended_validations": generate_validation_rules(state["t5_data"], state["disclosure_mapping"])
        }
        
        # Update requirements with data-driven scenarios
        if "extracted_requirements" not in state:
            state["extracted_requirements"] = {}
        
        state["extracted_requirements"]["data_driven_scenarios"] = state["structured_data_insights"]
        state["messages"].append(AIMessage(content="Completed structured data analysis"))
        
        return state
    
    return analyze_structured_data

def identify_quality_issues(t5_data: Dict, mapping_data: Dict) -> List[Dict]:
    """Identify potential data quality issues"""
    issues = []
    
    # Check T5 data quality
    for sheet_name, sheet_data in t5_data.get("sheets", {}).items():
        quality = sheet_data.get("quality_metrics", {})
        
        if quality.get("completeness", 1) < 0.95:
            issues.append({
                "type": "completeness",
                "source": f"T5 - {sheet_name}",
                "severity": "high" if quality.get("completeness", 1) < 0.8 else "medium",
                "description": f"Data completeness is {quality.get('completeness', 0)*100:.1f}%"
            })
        
        if quality.get("duplicate_percentage", 0) > 5:
            issues.append({
                "type": "duplicates",
                "source": f"T5 - {sheet_name}",
                "severity": "medium",
                "description": f"Found {quality.get('duplicate_percentage', 0):.1f}% duplicate rows"
            })
    
    return issues

def generate_validation_rules(t5_data: Dict, mapping_data: Dict) -> List[Dict]:
    """Generate validation rules based on data analysis"""
    rules = []
    
    # Generate rules for each column based on patterns
    for sheet_name, sheet_data in t5_data.get("sheets", {}).items():
        for col, analysis in sheet_data.get("column_analysis", {}).items():
            if analysis.get("patterns", {}).get("pattern_type") == "percentage":
                rules.append({
                    "field": col,
                    "rule": "value_between_0_and_100",
                    "description": f"Validate {col} is between 0 and 100"
                })
            
            if analysis.get("null_count", 0) == 0:
                rules.append({
                    "field": col,
                    "rule": "not_null",
                    "description": f"Validate {col} is never null"
                })
    
    return rules

def create_intelligent_test_generator(llm: ChatOpenAI):
    """Generate test cases with intelligent prioritization"""
    
    def generate_intelligent_tests(state: ConversationState) -> Dict:
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are a senior QA architect generating comprehensive test cases.
            
            GENERATION RULES:
            1. PRIORITIZATION (STRICT):
               - 60% from FSD requirements (PRIMARY SOURCE)
               - 25% from structured data insights (T5 & mappings)
               - 10% from edge cases and boundaries
               - 5% format reference from regression (STRUCTURE ONLY)
            
            2. TEST CASE STRUCTURE:
               - TC_ID: Unique identifier (TC_YYYY_MM_DD_XXX)
               - Test Name: Clear, actionable name
               - Objective: What is being tested
               - Priority: Critical/High/Medium/Low
               - Category: Functional/Data/Performance/Security/Integration
               - Prerequisites: Required setup
               - Test Steps: Detailed, numbered steps
               - Test Data: Specific data values to use
               - Expected Results: Clear pass criteria
               - Notes: Additional considerations
            
            3. COVERAGE REQUIREMENTS:
               - Every FSD requirement must have at least one test
               - Critical requirements need positive and negative tests
               - Data validations need boundary tests
               - Performance requirements need load tests
            
            4. QUALITY CRITERIA:
               - Tests must be atomic (test one thing)
               - Tests must be repeatable
               - Tests must have clear pass/fail criteria
               - Tests must be independent
            
            Generate diverse, comprehensive test cases that ensure complete coverage."""),
            ("human", """Requirements Analysis:
            {requirements}
            
            Structured Data Insights:
            {data_insights}
            
            Quality Issues to Test:
            {quality_issues}
            
            Regression Format (structure only):
            {regression_format}
            
            Generate comprehensive test cases with proper prioritization:""")
        ])
        
        # Prepare comprehensive input
        requirements_json = json.dumps(state["extracted_requirements"], indent=2)[:6000]
        data_insights = json.dumps(state.get("structured_data_insights", {}), indent=2)[:3000]
        quality_issues = identify_quality_issues(state["t5_data"], state["disclosure_mapping"])
        
        # Minimal regression reference
        regression_format = "TC Format: ID | Name | Steps | Expected Result" if state["regression_samples"] else "No format reference"
        
        chain = prompt | llm
        response = chain.invoke({
            "requirements": requirements_json,
            "data_insights": data_insights,
            "quality_issues": json.dumps(quality_issues, indent=2),
            "regression_format": regression_format
        })
        
        # Parse test cases
        test_cases = parse_test_cases_advanced(response.content if hasattr(response, 'content') else str(response))
        
        # Add metadata to each test case
        for i, tc in enumerate(test_cases):
            tc["generated_date"] = datetime.now().isoformat()
            tc["iteration"] = state.get("iteration", 0)
            tc["source_priority"] = determine_source_priority(tc, state)
        
        state["generated_test_cases"] = test_cases
        state["test_case_versions"].append(test_cases.copy())
        state["messages"].append(AIMessage(content=f"Generated {len(test_cases)} comprehensive test cases"))
        
        # Calculate quality metrics
        state["quality_metrics"] = calculate_quality_metrics(test_cases, state["extracted_requirements"])
        
        return state
    
    return generate_intelligent_tests

def parse_test_cases_advanced(text: str) -> List[Dict[str, Any]]:
    """Advanced parser for test cases with better structure extraction"""
    test_cases = []
    current_test = {}
    current_field = None
    
    lines = text.split('\n')
    
    field_mappings = {
        "tc_id": ["tc_id", "test case id", "id", "test_id"],
        "name": ["test name", "name", "title", "test case name"],
        "objective": ["objective", "purpose", "goal"],
        "priority": ["priority", "severity", "importance"],
        "category": ["category", "type", "test type"],
        "prerequisites": ["prerequisites", "preconditions", "setup"],
        "steps": ["test steps", "steps", "procedure"],
        "test_data": ["test data", "data", "input data"],
        "expected_results": ["expected results", "expected", "expected outcome"],
        "notes": ["notes", "comments", "additional info"]
    }
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line starts a new field
        line_lower = line.lower()
        field_found = False
        
        for field, keywords in field_mappings.items():
            if any(keyword in line_lower for keyword in keywords):
                # Save previous test case if complete
                if current_test and "tc_id" in current_test and "name" in current_test:
                    test_cases.append(current_test)
                    if field == "tc_id":  # Starting new test case
                        current_test = {}
                
                current_field = field
                # Extract value if on same line
                for keyword in keywords:
                    if keyword in line_lower:
                        value = line.split(':', 1)[-1].strip()
                        if value and value != line:
                            current_test[field] = value
                        break
                field_found = True
                break
        
        # If not a field header, add to current field
        if not field_found and current_field:
            if current_field in current_test:
                current_test[current_field] += f"\n{line}"
            else:
                current_test[current_field] = line
    
    # Add last test case
    if current_test and "tc_id" in current_test:
        test_cases.append(current_test)
    
    # Generate IDs if missing
    for i, tc in enumerate(test_cases):
        if "tc_id" not in tc or not tc["tc_id"]:
            tc["tc_id"] = f"TC_{datetime.now().strftime('%Y%m%d')}_{i+1:03d}"
    
    return test_cases

def determine_source_priority(test_case: Dict, state: ConversationState) -> str:
    """Determine which source primarily influenced the test case"""
    tc_text = json.dumps(test_case).lower()
    
    # Check for FSD-specific keywords
    fsd_keywords = ["business rule", "requirement", "validation", "workflow"]
    t5_keywords = ["data quality", "null value", "data type", "aggregation"]
    mapping_keywords = ["disclosure", "mapping", "relationship", "report"]
    
    fsd_score = sum(1 for kw in fsd_keywords if kw in tc_text)
    t5_score = sum(1 for kw in t5_keywords if kw in tc_text)
    mapping_score = sum(1 for kw in mapping_keywords if kw in tc_text)
    
    if fsd_score >= t5_score and fsd_score >= mapping_score:
        return "FSD"
    elif t5_score > mapping_score:
        return "T5"
    else:
        return "Mapping"

def calculate_quality_metrics(test_cases: List[Dict], requirements: Dict) -> Dict[str, Any]:
    """Calculate comprehensive quality metrics for generated test cases"""
    metrics = {
        "total_test_cases": len(test_cases),
        "requirement_coverage": {},
        "priority_distribution": {},
        "category_distribution": {},
        "completeness_score": 0,
        "quality_score": 0
    }
    
    # Count priorities and categories
    for tc in test_cases:
        priority = tc.get("priority", "Medium")
        category = tc.get("category", "Functional")
        
        metrics["priority_distribution"][priority] = metrics["priority_distribution"].get(priority, 0) + 1
        metrics["category_distribution"][category] = metrics["category_distribution"].get(category, 0) + 1
    
    # Calculate requirement coverage
    total_reqs = 0
    covered_reqs = 0
    
    for req_type, req_list in requirements.items():
        if isinstance(req_list, list):
            total_reqs += len(req_list)
            # Simple coverage check - in production, use more sophisticated matching
            covered = sum(1 for req in req_list if any(str(req)[:20] in str(tc) for tc in test_cases))
            covered_reqs += covered
            metrics["requirement_coverage"][req_type] = {
                "total": len(req_list),
                "covered": covered,
                "percentage": (covered / len(req_list) * 100) if req_list else 0
            }
    
    # Calculate scores
    metrics["completeness_score"] = (covered_reqs / total_reqs * 100) if total_reqs > 0 else 0
    
    # Quality score based on test case completeness
    complete_tests = sum(1 for tc in test_cases if all(
        field in tc for field in ["tc_id", "name", "steps", "expected_results"]
    ))
    metrics["quality_score"] = (complete_tests / len(test_cases) * 100) if test_cases else 0
    
    return metrics
